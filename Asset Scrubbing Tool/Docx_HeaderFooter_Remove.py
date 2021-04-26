#!/usr/bin/env python
# coding: utf-8

# In[ ]:


from docx import Document
import os
import docx
import re


# In[ ]:


def header_text_removal(Scrubbing_Location,Remove_Keywords,Scrubbing_docsFileList):
    print("scrubbing_Basedon_HeaderRemove_called")
       
    #functionForRemoveHeader
    def HeaderRemoveFromDoc():
        for File in Scrubbing_docsFileList:
            document = Document(File)
            section = document.sections[0]
            header = section.header
            
            for paragraph in header.paragraphs:
                header = document.sections[0].header
                allHeaderInHeader = header.paragraphs[0].text.split()
                for k in Remove_Keywords:
                    for v in allHeaderInHeader:
                        if k.lower() == v.lower():
                            allHeaderInHeader.remove(v)
                x = " ".join(allHeaderInHeader)
                header.paragraphs[0].text = x


            footer = section.footer
            for paragraph in footer.paragraphs:
                footer = document.sections[0].footer
                allHeaderInHeader = footer.paragraphs[0].text.split()
                for k in Remove_Keywords:
                    for v in allHeaderInHeader:
                        if k.lower() == v.lower():
                            allHeaderInHeader.remove(v)
                x = " ".join(allHeaderInHeader)
                footer.paragraphs[0].text = x


            document.save(File)

    HeaderRemoveFromDoc()
    
    local_docsFileList=[]
    del local_docsFileList

