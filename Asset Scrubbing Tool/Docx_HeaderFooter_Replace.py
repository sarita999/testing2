#!/usr/bin/env python
# coding: utf-8

# In[ ]:


from docx import Document
import os
import docx
import re


# In[ ]:


def header_text_Replace(Scrubbing_Location,Replace_List,Scrubbing_docsFileList):
    print("scrubbing_Basedon_HeaderReplace_called")
       
    #functionForRemoveHeader
    def HeaderReplaceFromDoc():
        for File in Scrubbing_docsFileList:
            document = Document(File)
            section = document.sections[0]

            header = section.header
            for paragraph in header.paragraphs:
                header = document.sections[0].header
                allHeaderInHeader = header.paragraphs[0].text.split() #it will print header
                print(allHeaderInHeader)
                for k1,v1 in Replace_List:
                    for v in allHeaderInHeader:
                        if k1.lower() == v.lower():
                            allHeaderInHeader.remove(v)
                            allHeaderInHeader.append(v1)
                    x1 = " ".join(allHeaderInHeader)
                    header.paragraphs[0].text = x1

            footer = section.footer
            for paragraph in footer.paragraphs:
                footer = document.sections[0].footer
                allHeaderInFooter = footer.paragraphs[0].text.split() #it will print header
                for k1,v1 in Replace_List:
                    for v in allHeaderInFooter:
                        if k1.lower() == v.lower():
                            allHeaderInFooter.remove(v)
                            allHeaderInFooter.append(v1)
                    x2 = " ".join(allHeaderInFooter)
                    footer.paragraphs[0].text = x2


            document.save(File)
                    
    HeaderReplaceFromDoc()
    
    local_docsFileList=[]
    del local_docsFileList
    
    


# In[ ]:




