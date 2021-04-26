#!/usr/bin/env python
# coding: utf-8

# In[ ]:


from docx import Document
import os
import docx
import re
import Docx_HeaderFooter_Remove
import FileName_Remove
import Docx_ReviewComment_Removal
import Docx_HeaderFooter_Replace
import FileName_Replace


# In[ ]:


def Scrubbing_Process(Location,Keywords,Flag,docsFileList):
    global Scrubbing_Location
    global Scrubbing_Keywords
    global Scrubbing_docsFileList
    
    Scrubbing_Location=Location
    Scrubbing_Keywords=Keywords
    Scrubbing_docsFileList=docsFileList
    Flag=Flag.lower()
    
    
    #if (Flag=="Replace")|(Flag=="replace")|(Flag=="REPLACE"):
    if (Flag == 'replace'):
        Scrubbing_Replace()
    elif (Flag == 'remove'):
        Scrubbing_Remove()
        


# In[ ]:


def Scrubbing_Replace():
    
    Replace_List=[]
    print("Replace-")
    Replace_Keywords=Scrubbing_Keywords.split(",")
    for i in Replace_Keywords:
        Single_Keyword=i.split(":")
        Replace_List.append(Single_Keyword)
    print(Replace_List)
    
    global dictOfReplacedInput
    dictOfReplacedInput= dict(x.split(':') for x in Scrubbing_Keywords.split(','))
    print(dictOfReplacedInput)
        
    ##Common Function to remove/replace for .doc file
    def docElementsRemove(docObj, regex, replace):
        for para in docObj.paragraphs:
            if regex.search(para.text):
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text
        for table in docObj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docElementsRemove(cell, regex, replace)


                                   
                    
    ### function to replace the specified text in doc file
    def replaceWordFromDoc():
        global textInWord
        global fullText
        global allWords
        global allWords1
        global allwords2
        fullText = []
        textInWord = []
        allWords = []
        allWords1 = []
        allwords2 = []
        for File in Scrubbing_docsFileList:
            try:
                doc = docx.Document(File)
                for para in doc.paragraphs:
                    fullText.append(para.text)
                    textInWord = '\n'.join(fullText).split()
                    allWords.extend(re.split(r'[->~\\\]\[!|#$%^&*();:,."/<?=+`{}\s]', para.text))
                    allWords = [x for x in allWords if x]
                for i,j in dictOfReplacedInput.items():
                    regex1 = re.compile(r'\b'+i+r'\b', flags=re.IGNORECASE)
                    replace1 = j
                    doc = Document(File)
                    docElementsRemove(doc, regex1, replace1)
                    doc.save(File)
            except:
                print("Couldn't open the file :",File)
                
    replaceWordFromDoc()
    Docx_HeaderFooter_Replace.header_text_Replace(Scrubbing_Location,Replace_List,Scrubbing_docsFileList)
    FileName_Replace.FileName_Replace(Scrubbing_Location,Scrubbing_Keywords,Scrubbing_docsFileList)

    

def Scrubbing_Remove():
    print("Remove Keywords-")
    Remove_Keywords=Scrubbing_Keywords.split(",")
    print(Remove_Keywords)
    
        
    #function for removing Element from docx
    def docElementsRemove(docObj, regex, replace):
        #Removing Data from paragraphs
        for para in docObj.paragraphs:
            if regex.search(para.text):
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text
        #Removing Data from Tables
        for table in docObj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docElementsRemove(cell, regex, replace) 
                    

    #function for removing Element from docx
    def removeWordFromDoc():
        global textInWord
        global fullText
        global allWords
        global allWords1
        global allwords2
        fullText = []
        textInWord = []
        allWords = []
        allWords1 = []
        allwords2 = []
        #Calling the docx files
        for File in Scrubbing_docsFileList:
            try:
                doc = docx.Document(File)
                #Accessing data in the docx files
                for para in doc.paragraphs:
                    fullText.append(para.text)
                    textInWord = '\n'.join(fullText).split()
                    allWords.extend(re.split(r'[->~\\\]\[!|#$%^&*();:,."/<?=+`{}\s]', para.text))
                    allWords = [x for x in allWords if x]              
                for l in Remove_Keywords:
                    regex1 = re.compile(r'\b'+l+r'\b', flags=re.IGNORECASE)
                    replace1 = r""
                    doc = Document(File)
                    docElementsRemove(doc, regex1, replace1)
                    doc.save(File)
            except:
                print("Couldn't open the file :",File)
                
    removeWordFromDoc()
    Docx_HeaderFooter_Remove.header_text_removal(Scrubbing_Location,Remove_Keywords,Scrubbing_docsFileList)
    #Docx_ReviewComment_Removal.reviewcommentsRemovalFromDoc(Scrubbing_Location,Remove_Keywords,Scrubbing_docsFileList)
    FileName_Remove.FileName_Removal(Scrubbing_Location,Scrubbing_Keywords,Scrubbing_docsFileList)



# In[ ]:




